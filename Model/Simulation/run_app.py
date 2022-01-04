import numpy as np
from skmultiflow.data import FileStream
from skmultiflow.neural_networks import PerceptronMask
from sklearn.metrics import accuracy_score
from Model.OFS import stability as st
import warnings
import matplotlib.pyplot as plt
from Model.OFS import fires as fires
from docx import Document
import timeit
import os
warnings.simplefilter(action='ignore', category=FutureWarning)

def apply_fires(df_name, tgt_index, epochs=1, batch_sizes = [25, 50, 75, 100], fractions = [0.1, 0.15, 0.2]):
    final_stab_lst = []
    final_acc_lst = []
    document = Document()

    for batch_size in batch_sizes:
        final_stab_lst_per_batch = []
        final_acc_lst_per_batch = []
        start = timeit.default_timer()
        for frac_selected_ftr in fractions:
            # Load data as scikit-multiflow FileStream
            # NOTE: FIRES accepts only numeric values. Please one-hot-encode or factorize string/char variables
            # Additionally, we suggest users to normalize all features, e.g. by using scikit-learn's MinMaxScaler()
            stream = FileStream(df_name, target_idx=tgt_index)
            stream.prepare_for_use()

            # Initial fit of the predictive model
            predictor = PerceptronMask()
            x, y = stream.next_sample(batch_size=batch_size)
            predictor.partial_fit(x, y, stream.target_values)

            # Initialize FIRES
            fires_model = fires.FIRES(n_total_ftr=stream.n_features,  # Total no. of features
                                target_values=stream.target_values,  # Unique target values (class labels)
                                mu_init=0,  # Initial importance parameter
                                sigma_init=1,  # Initial uncertainty parameter
                                penalty_s=0.01,
                                # Penalty factor for the uncertainty (corresponds to gamma_s in the paper)
                                penalty_r=0.01,
                                # Penalty factor for the regularization (corresponds to gamma_r in the paper)
                                epochs=epochs,
                                # No. of epochs that we use each batch of observations to update the parameters
                                lr_mu=0.01,  # Learning rate for the gradient update of the importance
                                lr_sigma=0.01,  # Learning rate for the gradient update of the uncertainty
                                scale_weights=True,  # If True, scale feature weights into the range [0,1]
                                model='probit')  # Name of the base model to compute the likelihood

            # Variables for calculating the average accuracy and stability per time step
            n_selected_ftr = round(frac_selected_ftr * stream.n_features)
            sum_acc, sum_stab, count_time_steps, stability_mat = 0, 0, 0, []
            stability_counter = 0
            start_window, end_window = 0, 9

            while stream.has_more_samples():
                # Load a new sample
                x, y = stream.next_sample(batch_size=batch_size)

                # Select features
                ftr_weights = fires_model.weigh_features(x, y)  # Get feature weights with FIRES
                ftr_selection = np.argsort(ftr_weights)[::-1][:n_selected_ftr]

                # Truncate x (retain only selected features, 'remove' all others, e.g. by replacing them with 0)
                x_reduced = np.zeros(x.shape)
                x_reduced[:, ftr_selection] = x[:, ftr_selection]

                # Prepare x to stability
                x_binary = np.zeros(stream.n_features)
                x_binary[ftr_selection] = 1
                stability_mat.append(x_binary)

                # Test
                y_pred = predictor.predict(x_reduced)
                acc_score = accuracy_score(y, y_pred)
                # print(acc_score)

                # Sum all the accuracy scores
                sum_acc = sum_acc + acc_score

                # Sum all the stabilty scores (shifting window = 10)
                if len(stability_mat) >= 10:
                    sum_stab = sum_stab + st.getStability(stability_mat[start_window:end_window])
                    # print(st.getStability(stability_mat[start_window:end_window]))
                    start_window += 1
                    end_window += 1
                    stability_counter += 1

                # Sum the time steps
                count_time_steps += 1

                # Train
                predictor.partial_fit(x_reduced, y)

            # Average accuracy  and stability
            avg_acc = sum_acc / count_time_steps
            avg_stab = sum_stab / (stability_counter)
            # print(f'avg acc score: {avg_acc}')
            # print(f'stability score: {avg_stab}')

            final_stab_lst.append(avg_stab)
            final_acc_lst.append(avg_acc)
            final_stab_lst_per_batch.append(avg_stab)
            final_acc_lst_per_batch.append(avg_acc)
            # Restart the FileStream
            stream.restart()

        plt.figure()
        plt.plot(fractions, final_acc_lst_per_batch, 'b', label='Accuracy over batch size of {}'.format(batch_size))
        plt.xlabel('Fraction')
        plt.ylabel('Accuracy')
        plt.legend()
        plt.savefig('plot.png')

        p = document.add_paragraph()
        r = p.add_run()
        r.add_text('Accuracy over batch size of {}'.format(batch_size))
        r.add_picture('plot.png')
        stop = timeit.default_timer()
        r.add_text('Runtime: {} seconds'.format(stop - start))
        os.remove('plot.png')

    # document.save('report.docx')

    print(f'Final avg acc score: {sum(final_acc_lst) / len(final_acc_lst)}')
    print(f'Final avg stab score: {sum(final_stab_lst) / len(final_stab_lst)}')

if __name__ == "__main__":
    apply_fires(df_name='/Users/samuelbenichou/Downloads/normalize/electricity_data.csv', tgt_index=0, epochs=1)